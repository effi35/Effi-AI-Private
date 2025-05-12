import os
import hashlib
import re
import magic  # לזיהוי סוגי קבצים
import chardet
import subprocess
from typing import Dict, List, Any, Tuple, Set, Optional
import importlib
import logging

logger = logging.getLogger(__name__)

class LanguageDetector:
    """מחלקה לזיהוי שפות תכנות"""
    
    def __init__(self):
        self.language_patterns = {
            "python": [r'\.py$', r'import\s+', r'from\s+\w+\s+import', r'def\s+\w+\s*\(', r'class\s+\w+\s*\('],
            "javascript": [r'\.js$', r'\.jsx$', r'function\s+\w+\s*\(', r'const\s+\w+\s*=', r'var\s+\w+\s*=', r'let\s+\w+\s*=', r'import\s+.*from', r'export\s+'],
            "typescript": [r'\.ts$', r'\.tsx$', r'interface\s+\w+', r'type\s+\w+', r'class\s+\w+'],
            "java": [r'\.java$', r'public\s+class', r'private\s+\w+', r'protected\s+\w+', r'package\s+\w+', r'import\s+\w+'],
            "c": [r'\.c$', r'\.h$', r'#include', r'int\s+main\s*\('],
            "cpp": [r'\.cpp$', r'\.hpp$', r'#include', r'namespace\s+\w+', r'template\s*<'],
            "csharp": [r'\.cs$', r'namespace\s+\w+', r'using\s+\w+', r'public\s+class'],
            "go": [r'\.go$', r'package\s+\w+', r'import\s+\(', r'func\s+\w+\s*\('],
            "ruby": [r'\.rb$', r'require\s+', r'def\s+\w+', r'class\s+\w+', r'module\s+\w+'],
            "php": [r'\.php$', r'\<\?php', r'function\s+\w+\s*\(', r'class\s+\w+'],
            "rust": [r'\.rs$', r'fn\s+\w+', r'struct\s+\w+', r'impl\s+', r'use\s+\w+'],
            "swift": [r'\.swift$', r'import\s+\w+', r'func\s+\w+', r'class\s+\w+', r'struct\s+\w+'],
            "kotlin": [r'\.kt$', r'fun\s+\w+', r'class\s+\w+', r'val\s+\w+', r'var\s+\w+'],
            "scala": [r'\.scala$', r'def\s+\w+', r'class\s+\w+', r'object\s+\w+', r'trait\s+\w+'],
            "html": [r'\.html$', r'\.htm$', r'\<html', r'\<head', r'\<body'],
            "css": [r'\.css$', r'\{', r'\.[\w-]+\s*\{', r'#[\w-]+\s*\{'],
            "xml": [r'\.xml$', r'\<\?xml', r'\<[a-zA-Z]+(\s+[\w-]+=\".*?\")*\s*\/?\>'],
            "json": [r'\.json$', r'\{\"', r'\[\"'],
            "yaml": [r'\.ya?ml$', r'[\w-]+\s*:\s*[\w\s-]+'],
            "dart": [r'\.dart$', r'void\s+main', r'class\s+\w+', r'import\s+'],
            "shell": [r'\.sh$', r'\#\!/bin/', r'function\s+\w+\s*\{'],
            "sql": [r'\.sql$', r'SELECT', r'INSERT', r'CREATE\s+TABLE']
        }
        
    def detect_language(self, file_path: str, content: Optional[str] = None) -> str:
        """זיהוי שפת התכנות לפי תוכן ושם הקובץ"""
        file_name = os.path.basename(file_path).lower()
        
        if content is None:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read(4096)  # קריאת 4KB ראשונים מספיקה לרוב
            except:
                content = ""
        
        # בדיקה לפי סיומת וביטויים רגולריים
        scores = {lang: 0 for lang in self.language_patterns}
        
        for lang, patterns in self.language_patterns.items():
            # בדיקת סיומת (ניקוד גבוה)
            for pattern in patterns:
                if pattern.startswith(r'\.') and re.search(pattern, file_name):
                    scores[lang] += 10
                elif content and re.search(pattern, content, re.IGNORECASE | re.MULTILINE):
                    scores[lang] += 1
        
        # בחירת השפה עם הניקוד הגבוה ביותר
        if max(scores.values()) > 0:
            return max(scores.items(), key=lambda x: x[1])[0]
        
        # אם לא זוהתה שפה, ננסה לפי סיומת
        ext = os.path.splitext(file_path)[1].lower()
        ext_to_lang = {
            '.py': 'python', '.js': 'javascript', '.ts': 'typescript', 
            '.java': 'java', '.c': 'c', '.cpp': 'cpp', '.cs': 'csharp',
            '.go': 'go', '.rb': 'ruby', '.php': 'php', '.rs': 'rust',
            '.swift': 'swift', '.kt': 'kotlin', '.scala': 'scala',
            '.html': 'html', '.css': 'css', '.xml': 'xml', '.json': 'json',
            '.yml': 'yaml', '.yaml': 'yaml', '.dart': 'dart', '.sh': 'shell',
            '.sql': 'sql', '.jsx': 'javascript', '.tsx': 'typescript'
        }
        
        return ext_to_lang.get(ext, 'unknown')

class MediaAnalyzer:
    """מחלקה לניתוח קבצי מדיה (תמונות, וידאו, אודיו)"""
    
    def __init__(self):
        self.supported_extensions = {
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.svg'],
            'video': ['.mp4', '.avi', '.mov', '.wmv', '.mkv', '.flv', '.webm'],
            'audio': ['.mp3', '.wav', '.ogg', '.flac', '.aac', '.m4a'],
        }
        
        # ניסיון לטעון תלויות לטיפול בתמונות
        self.image_support = False
        try:
            from PIL import Image
            self.Image = Image
            self.image_support = True
        except ImportError:
            logger.warning("לא ניתן לטעון תמיכה בתמונות (Pillow). ניתוח תמונות יוגבל.")
    
    def analyze_media_file(self, file_path: str) -> Dict[str, Any]:
        """ניתוח קובץ מדיה וחילוץ מטא-דאטה"""
        ext = os.path.splitext(file_path)[1].lower()
        media_type = self._get_media_type(ext)
        
        result = {
            "media_type": media_type,
            "metadata": {},
            "content_hash": self._calculate_hash(file_path)
        }
        
        if media_type == 'image' and self.image_support:
            try:
                img = self.Image.open(file_path)
                result["metadata"] = {
                    "format": img.format,
                    "mode": img.mode,
                    "width": img.width,
                    "height": img.height,
                    "exif": self._extract_exif(img)
                }
            except Exception as e:
                logger.error(f"שגיאה בניתוח תמונה {file_path}: {e}")
        
        # תוספת לניתוח קבצי וידאו ואודיו (באמצעות ffmpeg אם זמין)
        if media_type in ['video', 'audio']:
            try:
                result["metadata"] = self._extract_media_metadata(file_path, media_type)
            except Exception as e:
                logger.error(f"שגיאה בניתוח {media_type} {file_path}: {e}")
                
        return result
    
    def _get_media_type(self, ext: str) -> str:
        """קביעת סוג מדיה לפי סיומת"""
        for media_type, extensions in self.supported_extensions.items():
            if ext in extensions:
                return media_type
        return 'unknown'
    
    def _calculate_hash(self, file_path: str) -> str:
        """חישוב חתימת קובץ מדיה"""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception:
            return ""
    
    def _extract_exif(self, img) -> Dict:
        """חילוץ מטא-דאטה EXIF מתמונה"""
        exif_data = {}
        if hasattr(img, '_getexif') and img._getexif():
            for key, value in img._getexif().items():
                if isinstance(key, int) and isinstance(value, (str, int, float, bytes)):
                    exif_data[str(key)] = str(value)[:100]  # קיצור ערכים ארוכים
        return exif_data
    
    def _extract_media_metadata(self, file_path: str, media_type: str) -> Dict:
        """חילוץ מטא-דאטה מקבצי וידאו/אודיו באמצעות ffmpeg"""
        metadata = {}
        
        try:
            # בדיקה אם ffprobe (חלק מ-ffmpeg) זמין
            result = subprocess.run(
                ['ffprobe', '-version'], 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE,
                text=True
            )
            
            if result.returncode != 0:
                return {"error": "ffprobe לא זמין במערכת"}
            
            # הרצת ffprobe לחילוץ מטא-דאטה
            cmd = [
                'ffprobe', 
                '-v', 'quiet', 
                '-print_format', 'json', 
                '-show_format', 
                '-show_streams', 
                file_path
            ]
            
            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True
            )
            
            if result.returncode == 0 and result.stdout:
                import json
                data = json.loads(result.stdout)
                
                # חילוץ מטא-דאטה רלוונטי
                if 'format' in data:
                    metadata['format'] = data['format'].get('format_name', '')
                    metadata['duration'] = data['format'].get('duration', '0')
                    metadata['size'] = data['format'].get('size', '0')
                    metadata['bit_rate'] = data['format'].get('bit_rate', '0')
                
                if 'streams' in data and data['streams']:
                    if media_type == 'video' and any(s.get('codec_type') == 'video' for s in data['streams']):
                        video_stream = next(s for s in data['streams'] if s.get('codec_type') == 'video')
                        metadata['codec'] = video_stream.get('codec_name', '')
                        metadata['width'] = video_stream.get('width', 0)
                        metadata['height'] = video_stream.get('height', 0)
                        metadata['fps'] = eval(video_stream.get('r_frame_rate', '0/1'))
                    
                    if media_type in ['video', 'audio'] and any(s.get('codec_type') == 'audio' for s in data['streams']):
                        audio_stream = next(s for s in data['streams'] if s.get('codec_type') == 'audio')
                        metadata['audio_codec'] = audio_stream.get('codec_name', '')
                        metadata['channels'] = audio_stream.get('channels', 0)
                        metadata['sample_rate'] = audio_stream.get('sample_rate', '0')
        
        except Exception as e:
            metadata['error'] = f"שגיאה בחילוץ מטא-דאטה: {str(e)}"
        
        return metadata

class CodeAnalyzer:
    """מחלקה לניתוח מעמיק של קבצי קוד"""
    
    def __init__(self):
        self.language_detector = LanguageDetector()
        
        # הגדרת תבניות לזיהוי ייבוא וייצוא בשפות שונות
        self.import_patterns = {
            'python': [
                r'import\s+([\w\.]+)(?:\s+as\s+\w+)?',
                r'from\s+([\w\.]+)\s+import\s+(?:[\w\s,]+|\*)'
            ],
            'javascript': [
                r'import\s+(?:[\w\s,{}]*\s+from\s+)?[\'"]([^\'"]*)[\'"]\s*;?',
                r'require\s*\(\s*[\'"]([^\'"]*)[\'"]\s*\)\s*;?'
            ],
            'java': [
                r'import\s+([\w\.]+)(?:\.\*)?;'
            ],
            'rust': [
                r'use\s+([\w:]+)(?:::\*)?;'
            ],
            'go': [
                r'import\s+(?:\(\s*)?"([^"]*)"',
                r'import\s+"([^"]*)"'
            ]
        }
        
        self.export_patterns = {
            'python': [
                r'def\s+(\w+)\s*\(',
                r'class\s+(\w+)\s*(?:\([\w\s,]*\))?:'
            ],
            'javascript': [
                r'export\s+(?:default\s+)?(?:const|let|var|function|class)\s+(\w+)',
                r'(?:const|let|var|function|class)\s+(\w+).*\nexport\s+(?:\{[^}]*\}|default)',
                r'module\.exports\s*=',
                r'exports\.(\w+)\s*='
            ],
            'java': [
                r'public\s+(?:class|interface|enum)\s+(\w+)',
                r'public\s+(?:static\s+)?(?:final\s+)?(?:\w+)\s+(\w+)\s*\('
            ]
        }
        
        # ניסיון טעינת pygments לניתוח תחביר
        self.pygments_support = False
        try:
            import pygments
            from pygments import lexers, token
            self.pygments = pygments
            self.pygments_lexers = lexers
            self.pygments_token = token
            self.pygments_support = True
        except ImportError:
            logger.warning("לא ניתן לטעון תמיכה ב-Pygments. ניתוח תחביר מתקדם לא יהיה זמין.")
    
    def analyze_code(self, file_path: str, content: Optional[str] = None) -> Dict[str, Any]:
        """ניתוח מעמיק של קובץ קוד"""
        if content is None:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            except:
                logger.error(f"לא ניתן לקרוא את הקובץ {file_path}")
                return {
                    "language": "unknown",
                    "imports": [],
                    "exports": [],
                    "dependencies": [],
                    "functions": [],
                    "classes": [],
                    "keywords": []
                }
        
        # זיהוי שפה
        language = self.language_detector.detect_language(file_path, content)
        
        # ניתוח התוכן
        analysis = {
            "language": language,
            "imports": self._extract_imports(content, language),
            "exports": self._extract_exports(content, language),
            "functions": self._extract_functions(content, language),
            "classes": self._extract_classes(content, language),
            "keywords": self._extract_keywords(content),
            "dependencies": []
        }
        
        # ניתוח תחביר מתקדם באמצעות pygments (אם זמין)
        if self.pygments_support:
            try:
                tokens = self._tokenize_with_pygments(content, language)
                if tokens:
                    analysis["tokens_count"] = len(tokens)
                    analysis["syntax_elements"] = self._analyze_syntax_elements(tokens, language)
            except:
                pass
                
        # יצירת רשימת תלויות מהייבוא
        analysis["dependencies"] = self._convert_imports_to_dependencies(analysis["imports"], language)
                
        return analysis
    
    def _extract_imports(self, content: str, language: str) -> List[str]:
        """חילוץ ייבואים מקוד"""
        imports = []
        
        # אם יש תבניות לשפה זו
        if language in self.import_patterns:
            for pattern in self.import_patterns[language]:
                matches = re.finditer(pattern, content, re.MULTILINE)
                for match in matches:
                    if match.group(1) and match.group(1).strip():
                        imports.append(match.group(1).strip())
        else:
            # חיפוש כללי יותר אם אין תבניות ספציפיות
            lines = content.split('\n')
            for line in lines:
                if re.search(r'\b(import|require|include|use)\b', line, re.IGNORECASE):
                    # חיפוש מחרוזות אפשריות שמציינות ייבוא
                    matches = re.findall(r'[\'"]([^\'"]+)[\'"]', line)
                    imports.extend(matches)
        
        return list(set(imports))  # הסרת כפילויות
    
    def _extract_exports(self, content: str, language: str) -> List[str]:
        """חילוץ ייצוא מקוד"""
        exports = []
        
        # אם יש תבניות לשפה זו
        if language in self.export_patterns:
            for pattern in self.export_patterns[language]:
                matches = re.finditer(pattern, content, re.MULTILINE)
                for match in matches:
                    if match.groups() and match.group(1) and match.group(1).strip():
                        exports.append(match.group(1).strip())
        
        return list(set(exports))  # הסרת כפילויות
    
    def _extract_functions(self, content: str, language: str) -> List[str]:
        """חילוץ פונקציות מהקוד"""
        functions = []
        
        # התאמת תבנית לפי שפה
        if language == 'python':
            pattern = r'def\s+(\w+)\s*\('
        elif language in ['javascript', 'typescript']:
            pattern = r'function\s+(\w+)\s*\(|const\s+(\w+)\s*=\s*(?:async\s*)?\(.*\)\s*=>'
        elif language in ['java', 'csharp', 'cpp']:
            pattern = r'(?:public|private|protected)?\s*(?:static)?\s*\w+\s+(\w+)\s*\('
        elif language == 'go':
            pattern = r'func\s+(\w+)\s*\('
        elif language == 'rust':
            pattern = r'fn\s+(\w+)\s*\('
        else:
            # תבנית כללית
            pattern = r'\b(\w+)\s*\('
        
        matches = re.finditer(pattern, content, re.MULTILINE)
        for match in matches:
            for group in match.groups():
                if group:
                    functions.append(group)
                    break
        
        return list(set(functions))
    
    def _extract_classes(self, content: str, language: str) -> List[str]:
        """חילוץ מחלקות מהקוד"""
        classes = []
        
        # התאמת תבנית לפי שפה
        if language == 'python':
            pattern = r'class\s+(\w+)\s*(?:\(.*\))?:'
        elif language in ['javascript', 'typescript', 'java', 'csharp', 'cpp']:
            pattern = r'class\s+(\w+)'
        elif language == 'rust':
            pattern = r'struct\s+(\w+)|impl\s+.*for\s+(\w+)'
        else:
            # אין תבנית, מחזיר רשימה ריקה
            return []
        
        matches = re.finditer(pattern, content, re.MULTILINE)
        for match in matches:
            for group in match.groups():
                if group:
                    classes.append(group)
                    break
        
        return list(set(classes))
    
    def _extract_keywords(self, content: str) -> List[str]:
        """חילוץ מילות מפתח משמעותיות מהקוד"""
        # הסרת תווים מיוחדים והמרה לאותיות קטנות
        normalized = re.sub(r'[^\w\s]', ' ', content.lower())
        
        # פיצול למילים
        words = re.findall(r'\b\w{3,}\b', normalized)
        
        # ספירת שכיחות המילים
        from collections import Counter
        word_counts = Counter(words)
        
        # פילטור מילות עצירה נפוצות בקוד
        stop_words = {'the', 'and', 'this', 'that', 'for', 'from', 'with', 'have', 'not',
                      'int', 'str', 'bool', 'void', 'true', 'false', 'null', 'none'}
        
        # בחירת המילים השכיחות ביותר שאינן מילות עצירה
        keywords = [word for word, count in word_counts.most_common(30) 
                   if word not in stop_words and count > 1]
        
        return keywords
    
    def _tokenize_with_pygments(self, content: str, language: str):
        """פירוק לטוקנים באמצעות Pygments"""
        if not self.pygments_support:
            return None
            
        try:
            # ניסיון למצוא lexer מתאים
            try:
                lexer = self.pygments_lexers.get_lexer_by_name(language)
            except:
                try:
                    lexer = self.pygments_lexers.guess_lexer(content)
                except:
                    return None
            
            # טוקניזציה של התוכן
            tokens = list(lexer.get_tokens(content))
            return tokens
        except:
            return None
    
    def _analyze_syntax_elements(self, tokens, language: str) -> Dict[str, int]:
        """ניתוח אלמנטים תחביריים מרשימת טוקנים"""
        if not self.pygments_support:
            return {}
            
        element_counts = {
            'keywords': 0,
            'strings': 0,
            'comments': 0,
            'operators': 0,
            'identifiers': 0,
            'numbers': 0
        }
        
        for ttype, value in tokens:
            if ttype in self.pygments_token.Keyword:
                element_counts['keywords'] += 1
            elif ttype in self.pygments_token.String:
                element_counts['strings'] += 1
            elif ttype in self.pygments_token.Comment:
                element_counts['comments'] += 1
            elif ttype in self.pygments_token.Operator:
                element_counts['operators'] += 1
            elif ttype in self.pygments_token.Name:
                element_counts['identifiers'] += 1
            elif ttype in self.pygments_token.Number:
                element_counts['numbers'] += 1
        
        return element_counts
    
    def _convert_imports_to_dependencies(self, imports: List[str], language: str) -> List[str]:
        """המרת ייבואים לרשימת תלויות"""
        dependencies = []
        
        if language == 'python':
            # בפייתון, לרוב השם הראשון הוא שם החבילה
            for imp in imports:
                if '.' in imp:
                    package = imp.split('.')[0]
                else:
                    package = imp
                dependencies.append(package)
        elif language in ['javascript', 'typescript']:
            # בJS, הייבואים הם לרוב שמות מלאים של חבילות
            dependencies = imports
        elif language == 'java':
            # בJava, לרוב המילה הראשונה היא מזהה הארגון
            for imp in imports:
                if '.' in imp:
                    parts = imp.split('.')
                    if len(parts) >= 2:
                        org_name = '.'.join(parts[:2])
                        dependencies.append(org_name)
                    else:
                        dependencies.append(imp)
                else:
                    dependencies.append(imp)
        else:
            # ברירת מחדל, משתמשים בייבואים כמו שהם
            dependencies = imports
        
        return list(set(dependencies))

class DocumentAnalyzer:
    """מחלקה לניתוח מסמכים (טקסט, PDF, Office, וכו')"""
    
    def __init__(self):
        self.supported_extensions = {
            'text': ['.txt', '.md', '.rst', '.log', '.csv', '.tsv'],
            'pdf': ['.pdf'],
            'office': ['.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx'],
            'markup': ['.html', '.xml', '.json', '.yaml', '.yml', '.toml', '.ini', '.cfg']
        }
        
        # ניסיון לטעון תלויות לניתוח מסמכים
        self.pdf_support = False
        self.office_support = False
        
        try:
            # בדיקת תמיכה ב-PyPDF2
            import PyPDF2
            self.PyPDF2 = PyPDF2
            self.pdf_support = True
        except ImportError:
            logger.warning("לא ניתן לטעון תמיכה ב-PyPDF2. ניתוח מסמכי PDF יוגבל.")
        
        try:
            # בדיקת תמיכה ב-python-docx
            import docx
            self.docx = docx
            self.office_support = True
        except ImportError:
            logger.warning("לא ניתן לטעון תמיכה ב-python-docx. ניתוח מסמכי Office יוגבל.")
    
    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """ניתוח מסמך וחילוץ מידע"""
        ext = os.path.splitext(file_path)[1].lower()
        doc_type = self._get_document_type(ext)
        
        result = {
            "document_type": doc_type,
            "metadata": {},
            "content_preview": "",
            "content_hash": self._calculate_hash(file_path),
            "encoding": None,
            "keywords": []
        }
        
        # ניתוח לפי סוג המסמך
        if doc_type == 'text':
            self._analyze_text_file(file_path, result)
        elif doc_type == 'pdf' and self.pdf_support:
            self._analyze_pdf_file(file_path, result)
        elif doc_type == 'office' and self.office_support:
            self._analyze_office_file(file_path, result)
        elif doc_type == 'markup':
            self._analyze_markup_file(file_path, result)
        
        # חילוץ מילות מפתח מתצוגה מקדימה
        if result["content_preview"]:
            result["keywords"] = self._extract_keywords_from_text(result["content_preview"])
            
        return result
    
    def _get_document_type(self, ext: str) -> str:
        """קביעת סוג מסמך לפי סיומת"""
        for doc_type, extensions in self.supported_extensions.items():
            if ext in extensions:
                return doc_type
        return 'unknown'
    
    def _calculate_hash(self, file_path: str) -> str:
        """חישוב חתימת קובץ מסמך"""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception:
            return ""
    
    def _analyze_text_file(self, file_path: str, result: Dict[str, Any]) -> None:
        """ניתוח קובץ טקסט רגיל"""
        try:
            # זיהוי קידוד
            with open(file_path, 'rb') as f:
                raw_data = f.read(4096)  # קריאת התחלת הקובץ
            
            encoding_result = chardet.detect(raw_data)
            encoding = encoding_result['encoding'] or 'utf-8'
            result["encoding"] = encoding
            
            # קריאת תוכן לתצוגה מקדימה
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                content = f.read(2048)  # קריאת חלק ראשון
            
            result["content_preview"] = content
            result["metadata"]["lines_count"] = content.count('\n') + 1
            result["metadata"]["chars_count"] = len(content)
            
            # ניסיון לזהות פורמט
            if ext == '.csv':
                result["metadata"]["format"] = "CSV"
                result["metadata"]["delimiter"] = self._detect_csv_delimiter(content)
            elif ext == '.md':
                result["metadata"]["format"] = "Markdown"
            
        except Exception as e:
            logger.error(f"שגיאה בניתוח קובץ טקסט {file_path}: {e}")
    
    def _analyze_pdf_file(self, file_path: str, result: Dict[str, Any]) -> None:
        """ניתוח קובץ PDF"""
        if not self.pdf_support:
            return
            
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = self.PyPDF2.PdfReader(f)
                
                # מטא-דאטה
                result["metadata"]["pages_count"] = len(pdf_reader.pages)
                if pdf_reader.metadata:
                    result["metadata"]["author"] = pdf_reader.metadata.author
                    result["metadata"]["creator"] = pdf_reader.metadata.creator
                    result["metadata"]["producer"] = pdf_reader.metadata.producer
                    result["metadata"]["title"] = pdf_reader.metadata.title
                    result["metadata"]["subject"] = pdf_reader.metadata.subject
                
                # חילוץ טקסט מהעמוד הראשון
                if len(pdf_reader.pages) > 0:
                    first_page = pdf_reader.pages[0]
                    text = first_page.extract_text()
                    result["content_preview"] = text[:2048] if text else ""
                
        except Exception as e:
            logger.error(f"שגיאה בניתוח קובץ PDF {file_path}: {e}")
    
    def _analyze_office_file(self, file_path: str, result: Dict[str, Any]) -> None:
        """ניתוח קובץ Office (docx, xlsx, וכו')"""
        if not self.office_support:
            return
            
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.docx':
                doc = self.docx.Document(file_path)
                
                # מטא-דאטה
                result["metadata"]["paragraphs_count"] = len(doc.paragraphs)
                
                # חילוץ טקסט
                text = "\n".join([p.text for p in doc.paragraphs])
                result["content_preview"] = text[:2048] if text else ""
                
        except Exception as e:
            logger.error(f"שגיאה בניתוח קובץ Office {file_path}: {e}")
    
    def _analyze_markup_file(self, file_path: str, result: Dict[str, Any]) -> None:
        """ניתוח קובץ markup (html, xml, json, yaml, וכו')"""
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            # זיהוי קידוד
            with open(file_path, 'rb') as f:
                raw_data = f.read(4096)
            
            encoding_result = chardet.detect(raw_data)
            encoding = encoding_result['encoding'] or 'utf-8'
            result["encoding"] = encoding
            
            # קריאת תוכן
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                content = f.read(4096)
            
            result["content_preview"] = content[:2048]
            
            # ניתוח ספציפי לפי סוג
            if ext in ['.html', '.htm']:
                result["metadata"]["format"] = "HTML"
                # ניסיון לחלץ כותרת
                title_match = re.search(r'<title[^>]*>(.*?)</title>', content, re.IGNORECASE | re.DOTALL)
                if title_match:
                    result["metadata"]["title"] = title_match.group(1).strip()
            elif ext in ['.xml']:
                result["metadata"]["format"] = "XML"
            elif ext in ['.json']:
                result["metadata"]["format"] = "JSON"
                # בדיקת תקינות JSON
                try:
                    import json
                    json.loads(content)
                    result["metadata"]["valid"] = True
                except:
                    result["metadata"]["valid"] = False
            elif ext in ['.yaml', '.yml']:
                result["metadata"]["format"] = "YAML"
            
        except Exception as e:
            logger.error(f"שגיאה בניתוח קובץ markup {file_path}: {e}")
    
    def _detect_csv_delimiter(self, content: str) -> str:
        """זיהוי תו הפרדה בקובץ CSV"""
        possible_delimiters = [',', ';', '\t', '|']
        lines = content.split('\n')[:5]  # בדיקת 5 שורות ראשונות
        
        if not lines:
            return ','
        
        # בדיקת כל מפריד אפשרי וספירת ממוצע העמודות
        delimiter_counts = {}
        
        for delimiter in possible_delimiters:
            total_columns = 0
            for line in lines:
                if line.strip():
                    total_columns += len(line.split(delimiter))
            avg_columns = total_columns / len(lines) if lines else 0
            delimiter_counts[delimiter] = avg_columns
        
        # בחירת המפריד שנותן את מספר העמודות הגבוה ביותר
        return max(delimiter_counts.items(), key=lambda x: x[1])[0]
    
    def _extract_keywords_from_text(self, text: str) -> List[str]:
        """חילוץ מילות מפתח מטקסט"""
        # הסרת תווים מיוחדים והמרה לאותיות קטנות
        normalized = re.sub(r'[^\w\s]', ' ', text.lower())
        
        # פיצול למילים
        words = re.findall(r'\b\w{3,}\b', normalized)
        
        # ספירת שכיחות המילים
        from collections import Counter
        word_counts = Counter(words)
        
        # פילטור מילות עצירה נפוצות
        stop_words = {'the', 'and', 'this', 'that', 'for', 'from', 'with', 'have', 'not',
                      'are', 'you', 'your', 'can', 'will', 'all', 'any', 'they', 'there'}
        
        # בחירת המילים השכיחות ביותר שאינן מילות עצירה
        keywords = [word for word, count in word_counts.most_common(20) 
                   if word not in stop_words and count > 1]
        
        return keywords

class FileAnalyzer:
    """מחלקה מרכזית לניתוח קבצים - גרסה 2.0"""
    
    def __init__(self):
        self.initialized = False
        self.language_detector = LanguageDetector()
        self.media_analyzer = MediaAnalyzer()
        self.code_analyzer = CodeAnalyzer()
        self.document_analyzer = DocumentAnalyzer()
        
        # הגדרת רשימת סיומות בינאריות
        self.binary_extensions = {
            '.exe', '.dll', '.so', '.pyc', '.pyo', '.pyd', '.o', '.obj',
            '.class', '.jar', '.war', '.ear', '.zip', '.tar', '.gz', '.bz2',
            '.rar', '.7z', '.bin', '.dat', '.db', '.sqlite', '.mdb'
        }
    
    def initialize(self, config):
        """אתחול המנתח"""
        self.config = config
        self.initialized = True
        
        # יצירת לוגר
        self.logger = logging.getLogger(__name__)
        self.logger.info("מנתח הקבצים אותחל בהצלחה")
        
        return True
        
    def shutdown(self):
        """כיבוי המנתח"""
        self.initialized = False
        return True
        
    def analyze_file(self, file_path: str, rel_path: str) -> Dict[str, Any]:
        """ניתוח קובץ - מתודה מרכזית"""
        if not self.initialized:
            self.logger.error("מנתח הקבצים לא אותחל")
            return {"error": "מנתח הקבצים לא אותחל"}
        
        try:
            # המרת נתיב יחסי לשם קובץ בלבד אם צריך
            if '/' in rel_path or '\\' in rel_path:
                filename = os.path.basename(rel_path)
            else:
                filename = rel_path
                
            # חישוב hash לתוכן הקובץ
            content_hash = self._calculate_hash(file_path)
            
            # בדיקה האם הקובץ בינארי
            is_binary = self.is_binary_file(file_path)
            
            # בדיקת סוג הקובץ
            ext = os.path.splitext(file_path)[1].lower()
            file_type = self._get_file_type(ext)
            
            # תוצאות בסיסיות
            result = {
                "content_hash": content_hash,
                "is_binary": is_binary,
                "file_type": file_type,
                "file_size": os.path.getsize(file_path),
                "imports": [],
                "exports": [],
                "keywords": [],
                "relationships": []
            }
            
            # ניתוח מעמיק לפי סוג הקובץ
            if file_type == "code" and not is_binary:
                # ניתוח קוד
                code_analysis = self.code_analyzer.analyze_code(file_path)
                result["language"] = code_analysis["language"]
                result["imports"] = code_analysis["imports"]
                result["exports"] = code_analysis["exports"]
                result["functions"] = code_analysis["functions"]
                result["classes"] = code_analysis["classes"]
                result["keywords"] = code_analysis["keywords"]
                result["dependencies"] = code_analysis["dependencies"]
                if "syntax_elements" in code_analysis:
                    result["syntax_elements"] = code_analysis["syntax_elements"]
            
            elif file_type == "document" and not is_binary:
                # ניתוח מסמך
                doc_analysis = self.document_analyzer.analyze_document(file_path)
                result["document_type"] = doc_analysis["document_type"]
                result["metadata"] = doc_analysis["metadata"]
                result["encoding"] = doc_analysis["encoding"]
                result["keywords"] = doc_analysis["keywords"]
                result["content_preview"] = doc_analysis["content_preview"][:200]  # מקוצר
            
            elif file_type == "media":
                # ניתוח מדיה
                media_analysis = self.media_analyzer.analyze_media_file(file_path)
                result["media_type"] = media_analysis["media_type"]
                result["metadata"] = media_analysis["metadata"]
            
            # בניית מילות מפתח אם לא נוצרו
            if not result["keywords"] and not is_binary:
                result["keywords"] = self._extract_basic_keywords(file_path, filename)
            
            # ניתוח קשרים לקבצים אחרים
            result["relationships"] = self._analyze_relationships(file_path, result)
            
            return result
                
        except Exception as e:
            self.logger.error(f"שגיאה בניתוח קובץ {file_path}: {str(e)}")
            # החזרת תוצאות בסיסיות במקרה של שגיאה
            return {
                "content_hash": self._calculate_hash(file_path),
                "is_binary": self.is_binary_file(file_path),
                "file_type": self._get_file_type(os.path.splitext(file_path)[1].lower()),
                "error": str(e),
                "imports": [],
                "exports": [],
                "keywords": []
            }
    
    def _calculate_hash(self, file_path: str) -> str:
        """חישוב חתימת הקובץ"""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                return hashlib.md5(content).hexdigest()
        except:
            return ""
    
    def _get_file_type(self, ext: str) -> str:
        """קביעת סוג הקובץ לפי סיומת"""
        code_extensions = {
            ".py", ".js", ".java", ".c", ".cpp", ".cs", ".php", ".html", 
            ".css", ".xml", ".json", ".md", ".h", ".hpp", ".jsx", ".ts", 
            ".tsx", ".go", ".rb", ".rs", ".swift", ".kt", ".scala", ".sh", 
            ".bat", ".ps1", ".sql", ".yml", ".yaml", ".toml", ".dart", ".r"
        }
        
        markup_extensions = {
            ".html", ".xml", ".md", ".json", ".yml", ".yaml", ".toml", ".ini", ".config"
        }
        
        document_extensions = {
            ".txt", ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", 
            ".odt", ".ods", ".odp", ".rtf", ".csv", ".tsv", ".rst"
        }
        
        media_extensions = {
            ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp", ".svg",
            ".mp4", ".avi", ".mov", ".wmv", ".mkv", ".flv", ".webm",
            ".mp3", ".wav", ".ogg", ".flac", ".aac", ".m4a"
        }
        
        if ext in code_extensions:
            return "code"
        elif ext in document_extensions:
            return "document"
        elif ext in media_extensions:
            return "media"
        elif ext in markup_extensions:
            return "markup"
        else:
            return "unknown"
    
    def is_binary_file(self, file_path: str) -> bool:
        """בדיקה האם הקובץ בינארי"""
        # בדיקה לפי סיומת
        ext = os.path.splitext(file_path)[1].lower()
        if ext in self.binary_extensions:
            return True
        
        # בדיקה לפי תוכן
        try:
            with open(file_path, 'tr', encoding='utf-8') as f:
                f.read(1024)
            return False
        except:
            return True
    
    def _extract_basic_keywords(self, file_path: str, filename: str) -> List[str]:
        """חילוץ בסיסי של מילות מפתח מהקובץ ושמו"""
        keywords = []
        
        # הוספת מילים משם הקובץ
        name_without_ext = os.path.splitext(filename)[0]
        name_words = re.findall(r'\b\w{3,}\b', name_without_ext.lower())
        keywords.extend(name_words)
        
        # ניסיון לקרוא את התחלת הקובץ
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(1024)
                content_words = re.findall(r'\b\w{3,}\b', content.lower())
                
                # ספירת שכיחות המילים
                from collections import Counter
                word_counts = Counter(content_words)
                
                # הוספת המילים השכיחות ביותר
                for word, count in word_counts.most_common(10):
                    if count > 1 and word not in keywords:
                        keywords.append(word)
        except:
            pass
            
        return keywords
    
    def _analyze_relationships(self, file_path: str, analysis_results: Dict[str, Any]) -> List[Dict[str, Any]]:
        """ניתוח הקשרים של הקובץ לקבצים אחרים"""
        relationships = []
        
        # קשרים לפי ייבוא
        if "imports" in analysis_results and analysis_results["imports"]:
            for imp in analysis_results["imports"]:
                relationships.append({
                    "type": "import",
                    "target": imp,
                    "weight": 1.0
                })
        
        # קשרים לפי תלויות
        if "dependencies" in analysis_results and analysis_results["dependencies"]:
            for dep in analysis_results["dependencies"]:
                relationships.append({
                    "type": "dependency",
                    "target": dep,
                    "weight": 0.8
                })
        
        return relationships